# Description: This is a module containing a python library that
#              makes it easier to create word documents using python

#####################################################################
##                              Imports                            ##
#####################################################################
from docx import Document
from docx.text.paragraph import Paragraph
import pandas as pd
from docx import styles
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


#####################################################################
##                              Classes                            ##
#####################################################################

class word_doc:
    '''
    Example Usage:
    from ds_modules_101.Utilities import word_doc
    from ds_modules_101.Data import f_word_doc
    from docx.shared import RGBColor
    import pandas as pd

    # First create a word document to be a template
    # Get the template location
    template_path = f_word_doc

    # Create a dataframe
    df = pd.DataFrame(data=[[1,2,3,4],[10,20,30,40]],index=['bought','sold'],columns=['01','02','03','04'])

    # Read the word document
    my_word_doc = word_doc(template_path)

    # These are the colors we will use
    rgb_color_default = RGBColor(44,48,138)
    rgb_color_green = RGBColor(0,176,80)
    rgb_color_red = RGBColor(255,0,0)

    # REPLACING A VALUE IN THE DOCUMENT
    my_word_doc.replace_term_in_entire_doc('_title_','My Calendar For NOV',rgb_color=rgb_color_default,verbose=True)

    # REPLACING VALUES IN THE DOCUMENT FROM A DATAFRAME. For each cell in the dataframe, make the replacement and color
    for row in df.index:
        for col in df.columns:
            # We are looking for "_rowcol_" in the document to do the replacement
            handle = '_'+str(row)+str(col)+'_'

            # Get the relevant value for that cell
            value = df.loc[row,col]

            # Identify if we need to colour it in green or red
            rgb_color = rgb_color_default
            if value < 20:
                rgb_color = rgb_color_red
            else:
                rgb_color = rgb_color_green

            my_word_doc.replace_term_in_entire_doc(handle, str(value), rgb_color=rgb_color,
                                                   verbose=True)

    styles = my_word_doc.get_table_styles()

    print(styles)

    # Printing a table. The table style must already exist in the document
    my_word_doc.table_from_dataframe(df=df,include_index=True,style='Grid Table 1 Light Accent 1')

    # Color a particular word in a table to grey
    my_word_doc.color_term_in_table(my_word_doc.document.tables[-1],'03','grey')

    my_word_doc.save('word_doc_output.docx')
    '''

    def __init__(self, document_path):
        '''
        dfa
        '''
        self.document = Document(document_path)

    def replace_term_in_body(self, old_term, new_term, verbose=False, justification='center', rgb_color=None,
                             bold=None):
        i = 0
        for paragraph in self.document.paragraphs:
            if old_term in paragraph.text:
                paragraph.text = paragraph.text.replace(old_term, new_term)
                if justification.lower() == 'center':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif justification.lower() == 'left':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif justification.lower() == 'right':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                i = i + 1

        if verbose:
            print('{} occurences of {} replaced with {} in text body'.format(i, old_term, new_term))

    def replace_term_in_table_cells(self, table, old_term, new_term, justification='center', rgb_color=None, bold=None):
        i = 0
        n_cols = len(table.columns)
        n_rows = len(table.rows)
        for i_col in range(n_cols):
            for i_row in range(n_rows):
                try:
                    for paragraph in table.cell(i_row, i_col).paragraphs:
                        if old_term in paragraph.text:
                            paragraph.text = paragraph.text.replace(old_term, new_term)
                            if justification.lower() == 'center':
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif justification.lower() == 'left':
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            elif justification.lower() == 'right':
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                            for run in paragraph.runs:
                                if bold is not None:
                                    run.font.bold = True
                                if rgb_color is not None:
                                    run.font.color.rgb = rgb_color  # RGBColor(0x42, 0x24, 0xE9)

                            i = i + 1

                #                     if old_term in table.cell(i_row,i_col).text:
                #                         style = table.cell(i_row,i_col).paragraphs[0].style
                #                         table.cell(i_row,i_col).text = table.cell(i_row,i_col).text.replace(old_term,new_term)
                #                         table.cell(i_row,i_col).paragraphs[0].style = style
                #                         if justification.lower() == 'center':
                #                             table.cell(i_row,i_col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                #                         elif justification.lower() == 'left':
                #                             table.cell(i_row,i_col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                #                         elif justification.lower() == 'right':
                #                             table.cell(i_row,i_col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                #                         for run in table.cell(i_row,i_col).paragraphs[0].runs:
                #                             run.font.bold = True
                #                             run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

                #                             i = i+1
                except:
                    continue

        return i

    def replace_term_in_table_columns(self, table, old_term, new_term, justification='center', rgb_color=None,
                                      bold=None):
        i = 0
        n_cols = len(table.columns)
        for col in table.row_cells(0):
            if old_term in col.text:
                col.text = col.text.replace(old_term, new_term)
                if justification.lower() == 'center':
                    col.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif justification.lower() == 'left':
                    col.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif justification.lower() == 'right':
                    col.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                i = i + 1

        return i

    def replace_term_in_table_rows(self, table, old_term, new_term, justification='center', rgb_color=None, bold=None):
        i = 0
        n_cols = len(table.rows)
        for row in table.column_cells(0):
            if old_term in row.text:
                row.text = row.text.replace(old_term, new_term)

                if justification.lower() == 'center':
                    row.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif justification.lower() == 'left':
                    row.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif justification.lower() == 'right':
                    row.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                i = i + 1

        return i

    def replace_term_in_table(self, table, old_term, new_term, include_cells=True, include_cols=False,
                              include_rows=False,
                              justification='center', rgb_color=None, bold=None):
        i = 0

        if include_cells:
            i = i + self.replace_term_in_table_cells(table, old_term, new_term, justification=justification,
                                                     rgb_color=rgb_color, bold=bold)
        if include_cols:
            i = i + self.replace_term_in_table_columns(table, old_term, new_term, justification=justification,
                                                       rgb_color=rgb_color, bold=bold)
        if include_rows:
            i = i + self.replace_term_in_table_rows(table, old_term, new_term, justification=justification,
                                                    rgb_color=rgb_color, bold=bold)

        return i

    def replace_term_in_entire_doc(self, old_term, new_term, verbose=False, rgb_color=None, bold=None):
        i_body = self.replace_term_in_body(old_term, new_term, verbose=verbose, rgb_color=rgb_color, bold=bold)

        i_tables = 0
        for table in self.document.tables:
            i_tables = i_tables + self.replace_term_in_table(table, old_term, new_term, include_cells=True,
                                                             include_cols=True, include_rows=True,
                                                             rgb_color=rgb_color, bold=bold)

        if verbose:
            print('{} occurences of {} replace by {} in tables'.format(i_tables, old_term, new_term))

    def update_document_title(self, title):
        self.document.paragraphs[0].text = title

    def get_table_styles(self):
        return [style.name for style in self.document.styles if
                      type(style) == type(self.document.styles['Normal Table'])]

    def get_styles(self):
        return [style.name for style in self.document.styles]

    def add_table(self, df, include_index=False, style=None):
        if style is None:
            styles = [style.name for style in self.document.styles if
                      type(style) == type(self.document.styles['Normal Table'])]

            if 'Plain Table 4' in styles:
                style = self.document.styles['Plain Table 4']
            else:
                style = self.document.styles['Normal Table']

        if include_index:
            self.document.add_table(df.shape[0] + 1, df.shape[1] + 1)
        else:
            self.document.add_table(df.shape[0], df.shape[1])

        table = self.document.tables[-1]
        table.style = style
        return table

    def table_from_dataframe(self, df, create_table=True, include_index=False, table=None, style=None):
        if (not create_table) and (table is None):
            raise Exception('An existing table needs to be specified if create table is false')

        if create_table:
            table = self.add_table(df, include_index=include_index, style=style)

        if include_index:
            for idx, col in enumerate(table.row_cells(0)):
                col.text = list([''] + list(df.columns))[idx]

            for idx, row in enumerate(table.column_cells(0)):
                row.text = list([''] + list(df.index))[idx]

            n_cols = len(table.columns)
            n_rows = len(table.rows)
            for i_col in range(1, n_cols):
                for i_row in range(1, n_rows):
                    table.cell(i_row, i_col).text = str(df.iloc[i_row - 1, i_col - 1])

        else:
            for idx, col in enumerate(table.row_cells(0)):
                col.text = list(df.columns)[idx]

            n_cols = len(table.columns)
            n_rows = len(table.rows)
            n_start = 0
            for i_col in range(0, n_cols):
                for i_row in range(1, n_rows):
                    table.cell(i_row, i_col).text = str(df.iloc[i_row - 1, i_col - 1])

    def color_term_in_table(self, table, term, color=RGBColor(159, 160, 166), isInterval=False):
        if type(color) == type('string'):
            if color.lower() == 'red':
                color = RGBColor(242, 19, 19)
            elif color.lower() in ('grey', 'gray'):
                color = RGBColor(159, 160, 166)
            elif color.lower() in ('green'):
                color = RGBColor(19, 242, 30)
        for row in range(len(table.rows)):
            for cell in table.row_cells(row):
                if type(term) == type('string'):
                    if cell.text.lower() == term.lower():
                        run = cell.paragraphs[0].runs[0]
                        run.font.color.rgb = color
                elif isInterval:
                    try:
                        if (float(cell.text) >= term[0]) & (float(cell.text) <= term[1]):
                            run = cell.paragraphs[0].runs[0]
                            run.font.color.rgb = color
                    except:
                        continue
                else:
                    if str(cell.text) == str(term):
                        run = cell.paragraphs[0].runs[0]
                        run.font.color.rgb = color

    def get_styles(self):
        return [style for style in self.document.styles]

    def save(self, path):
        self.document.save(path)

if __name__ == '__main__':

    import os
    import sys
    import pandas as pd

    # First create a word document to be a template
    # Get the template location
    current_dir = '/'.join(sys.path[0].split('/')[:-1])  # sys.path[0]
    data_dir = os.path.join(current_dir, 'Data', 'Docs')
    template_path = os.path.join(data_dir, 'Word_Doc.docx')

    # Create a dataframe
    df = pd.DataFrame(data=[[1,2,3,4],[10,20,30,40]],index=['bought','sold'],columns=['01','02','03','04'])

    # Read the word document
    my_word_doc = word_doc(template_path)

    # These are the colors we will use
    rgb_color_default = RGBColor(44,48,138)
    rgb_color_green = RGBColor(0,176,80)
    rgb_color_red = RGBColor(255,0,0)

    # REPLACING A VALUE IN THE DOCUMENT
    my_word_doc.replace_term_in_entire_doc('_title_','My Calendar For NOV',rgb_color=rgb_color_default,verbose=True)

    # REPLACING VALUES IN THE DOCUMENT FROM A DATAFRAME. For each cell in the dataframe, make the replacement and color
    for row in df.index:
        for col in df.columns:
            # We are looking for "_rowcol_" in the document to do the replacement
            handle = '_'+str(row)+str(col)+'_'

            # Get the relevant value for that cell
            value = df.loc[row,col]

            # Identify if we need to colour it in green or red
            rgb_color = rgb_color_default
            if value < 20:
                rgb_color = rgb_color_red
            else:
                rgb_color = rgb_color_green

            my_word_doc.replace_term_in_entire_doc(handle, str(value), rgb_color=rgb_color,
                                                   verbose=True)

    styles = my_word_doc.get_table_styles()

    print(styles)

    # Printing a table. The table style must already exist in the document
    my_word_doc.table_from_dataframe(df=df,include_index=True,style='Grid Table 1 Light Accent 1')

    # Color a particular word in a table to grey
    my_word_doc.color_term_in_table(my_word_doc.document.tables[-1],'03','grey')

    my_word_doc.save('word_doc_output.docx')




